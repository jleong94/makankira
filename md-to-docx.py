#!/usr/bin/env python3
"""
md-to-docx.py - Convert the MakanKira Markdown spec to a .docx deliverable.

The Markdown file (meal-order-payment-app-document.md) is the canonical editing
source. After editing it, run this script to regenerate the .docx. No pandoc
required - it uses python-docx only.

Usage:
    python md-to-docx.py
        Converts meal-order-payment-app-document.md -> .docx in this folder.
    python md-to-docx.py input.md
        Converts input.md -> input.docx (same basename).
    python md-to-docx.py input.md output.docx
        Explicit input and output paths.

Requires: python-docx   (pip install python-docx)

Supported Markdown: ATX headings (#..####), pipe tables (with alignment and
escaped \\| ), fenced code blocks (```), bullet lists (- , nested with indent),
numbered lists (literal "N."), and inline **bold** / `code`.
"""
import sys, os, re

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx is required. Install it with:  pip install python-docx")

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_MD = os.path.join(SCRIPT_DIR, "meal-order-payment-app-document.md")

if len(sys.argv) >= 3:
    IN, OUT = sys.argv[1], sys.argv[2]
elif len(sys.argv) == 2:
    IN = sys.argv[1]
    OUT = os.path.splitext(IN)[0] + ".docx"
else:
    IN = DEFAULT_MD
    OUT = os.path.splitext(DEFAULT_MD)[0] + ".docx"

if not os.path.exists(IN):
    sys.exit("Input Markdown not found: " + IN)

with open(IN, "r", encoding="utf-8") as f:
    lines = f.read().split("\n")

doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
counts = {"headings": 0, "tables": 0, "code": 0}


def unescape(t):
    return re.sub(r"\\([\\`*_{}\[\]()#+\-.!|])", r"\1", t)


def add_runs(p, text):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            p.add_run(unescape(text[pos:m.start()]))
        tok = m.group(0)
        if tok.startswith("**"):
            r = p.add_run(unescape(tok[2:-2])); r.bold = True
        else:
            r = p.add_run(tok[1:-1].replace("\\|", "|")); r.font.name = "Consolas"; r.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        p.add_run(unescape(text[pos:]))


def split_row(line):
    s = line.strip()
    if s.startswith("|"): s = s[1:]
    if s.endswith("|"): s = s[:-1]
    return [c.strip() for c in re.split(r"(?<!\\)\|", s)]


def is_sep(line):
    if "|" not in line and "-" not in line:
        return False
    cells = split_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{1,}:?", c.strip() or "") for c in cells)


def aligns(sep):
    out = []
    for c in sep:
        c = c.strip(); l = c.startswith(":"); r = c.endswith(":")
        out.append(WD_ALIGN_PARAGRAPH.CENTER if (l and r) else WD_ALIGN_PARAGRAPH.RIGHT if r else WD_ALIGN_PARAGRAPH.LEFT)
    return out


def shade(el, fill):
    pr = el.get_or_add_tcPr() if el.tag.endswith("}tc") else el.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill); pr.append(shd)


def add_code(code_lines):
    counts["code"] += 1
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(4); pf.space_after = Pt(4); pf.left_indent = Pt(6)
    shade(p._p, "F2F2F2")
    for i, ln in enumerate(code_lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(ln); r.font.name = "Consolas"; r.font.size = Pt(9)


def add_table(rows):
    counts["tables"] += 1
    header = split_row(rows[0]); al = aligns(split_row(rows[1])); body = [split_row(r) for r in rows[2:]]
    ncol = len(header); t = doc.add_table(rows=1, cols=ncol); t.style = "Table Grid"
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]; para = cell.paragraphs[0]
        if j < len(al): para.alignment = al[j]
        add_runs(para, h.replace("\\|", "|"))
        for run in para.runs: run.bold = True
        shade(cell._tc, "D9E2F3")
    for brow in body:
        cells = t.add_row().cells
        for j in range(ncol):
            para = cells[j].paragraphs[0]
            if j < len(al): para.alignment = al[j]
            add_runs(para, (brow[j] if j < len(brow) else "").replace("\\|", "|"))
    doc.add_paragraph()


def add_head(text, level):
    counts["headings"] += 1
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text).replace("`", "")
    h = doc.add_heading(text, 0 if level == 1 else level - 1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


buf = []


def flush():
    global buf
    if buf:
        txt = " ".join(buf).strip()
        if txt:
            add_runs(doc.add_paragraph(), txt)
        buf = []


i, n = 0, len(lines)
while i < n:
    line = lines[i]; s = line.strip()
    if s.startswith("```"):
        flush(); code = []; i += 1
        while i < n and not lines[i].strip().startswith("```"):
            code.append(lines[i]); i += 1
        i += 1; add_code(code); continue
    m = re.match(r"^(#{1,6})\s+(.*)$", line)
    if m:
        flush(); add_head(m.group(2).strip(), len(m.group(1))); i += 1; continue
    if s.startswith("|") and i + 1 < n and is_sep(lines[i + 1]):
        flush(); tbl = [lines[i], lines[i + 1]]; i += 2
        while i < n and lines[i].strip().startswith("|"):
            tbl.append(lines[i]); i += 1
        add_table(tbl); continue
    if s == "":
        flush(); i += 1; continue
    mb = re.match(r"^(\s*)[-*]\s+(.*)$", line)
    if mb:
        flush(); p = doc.add_paragraph(style="List Bullet")
        if len(mb.group(1)) >= 2:
            p.paragraph_format.left_indent = Inches(0.75)
        add_runs(p, mb.group(2)); i += 1; continue
    mn = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
    if mn:
        flush(); p = doc.add_paragraph(); pf = p.paragraph_format
        pf.left_indent = Inches(0.5 if len(mn.group(1)) >= 2 else 0.25); pf.first_line_indent = Inches(-0.25)
        p.add_run(mn.group(2) + ". "); add_runs(p, mn.group(3)); i += 1; continue
    buf.append(s); i += 1
flush()

doc.save(OUT)
chk = Document(OUT)
print("Converted:", os.path.basename(IN), "->", os.path.basename(OUT))
print("  headings=%d tables=%d code_blocks=%d | reopened ok: paragraphs=%d tables=%d"
      % (counts["headings"], counts["tables"], counts["code"], len(chk.paragraphs), len(chk.tables)))
